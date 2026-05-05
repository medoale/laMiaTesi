"""
Genera vulnRadar_spiegazione.docx — documento Word in italiano
con la spiegazione logica del programma per il professore.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def p(text, bold=False, italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    return para


def bullet(text):
    doc.add_paragraph(text, style='List Bullet')


# =============================================================================
# TITOLO
# =============================================================================
title = doc.add_heading('vulnRadar — Sistema predittivo per il monitoraggio di repository GitHub a rischio vulnerabilità', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p('Documento di spiegazione tecnica e logica del programma', italic=True).alignment = WD_ALIGN_PARAGRAPH.CENTER

# =============================================================================
# 1. OBIETTIVO
# =============================================================================
h1('1. Obiettivo del programma')
p(
    'vulnRadar è un sistema automatico che, ogni giorno, individua un insieme di '
    'repository GitHub considerati "ad alto rischio" — ovvero più probabilmente '
    'oggetto di una futura vulnerabilità (CVE) pubblicata. L\'idea di fondo è '
    'costruire un osservatorio predittivo: selezioniamo i repository sospetti oggi, '
    'li archiviamo in un database, e ad ogni esecuzione successiva verifichiamo se '
    'qualcuno di quelli che avevamo già selezionato in passato è effettivamente '
    'stato menzionato in un nuovo CVE pubblicato sul database ufficiale NVD.'
)
p(
    'Il sistema è organizzato in tre task indipendenti che girano in parallelo, '
    'ciascuno con un proprio criterio di selezione, e in un quarto modulo di '
    'cross-reference che esegue il confronto storico con i CVE.'
)

# =============================================================================
# 2. ARCHITETTURA
# =============================================================================
h1('2. Architettura generale')
p(
    'Il programma è scritto in Python ed è composto da nove moduli con '
    'responsabilità ben separate:'
)
bullet('main.py — entry point, orchestra l\'esecuzione parallela e la persistenza')
bullet('config.py — parametri configurabili e lettura di .CVEfixes.ini per i token')
bullet('github_client.py — client HTTP thread-safe verso GitHub con gestione del rate limit')
bullet('nvd_client.py — client robusto verso il database NVD con finestratura, paginazione e gestione dei fallimenti parziali')
bullet('database.py — definizione dello schema SQLite e funzioni di insert/lettura')
bullet('task_official.py — Task 1, basato sul database ufficiale NVD')
bullet('task_hot.py — Task 2, basato su segnali di attività di sicurezza e silent patch')
bullet('task_talkers.py — Task 3, basato sui repository più "rumorosi"')
bullet('cve_matcher.py — modulo di confronto con i CVE pubblicati successivamente')

p(
    'I tre task vengono eseguiti contemporaneamente tramite ThreadPoolExecutor con '
    'tre worker: questo significa che le chiamate API verso GitHub e NVD avvengono '
    'in parallelo, riducendo il tempo totale di esecuzione e simulando il '
    'comportamento di un osservatorio realmente reattivo. Il client GitHub è '
    'thread-safe e serializza le pause causate dai limiti di velocità tramite un '
    'lock condiviso.'
)

# =============================================================================
# 3. I TRE CRITERI DI SELEZIONE
# =============================================================================
h1('3. I tre criteri di selezione')

# ----- Task 1 -----
h2('3.1 — Task "Official": vendor e prodotti più colpiti su NVD')
p(
    'La fonte è l\'API ufficiale del National Vulnerability Database (NVD) del NIST. '
    'Il task scarica tutti i CVE pubblicati negli ultimi 30 giorni '
    '(parametro NVD_LOOKBACK_DAYS).'
)
p(
    'Per ogni CVE, NVD fornisce le configurazioni dei prodotti vulnerabili tramite '
    'stringhe CPE (Common Platform Enumeration) nel formato:'
)
code('cpe:2.3:a:apache:struts:2.5.0:...')
p(
    'I campi rilevanti sono il quarto (vendor, in questo esempio "apache") e il '
    'quinto (product, "struts"). Il task estrae da ogni CVE l\'insieme delle coppie '
    '(vendor, product) — un CVE che lista cinque versioni dello stesso prodotto '
    'conta una sola volta per quella coppia, evitando il doppio conteggio. Le '
    'coppie sono poi ordinate per frequenza di apparizione nei CVE recenti.'
)
p(
    'Per ogni coppia in cima alla classifica, il sistema cerca di risolvere '
    'l\'effettivo repository GitHub che ospita quel prodotto, applicando in '
    'sequenza tre strategie:'
)
bullet('lookup diretto su /repos/{vendor}/{product}')
bullet('lookup su /repos/{vendor_mappato}/{product} per i casi in cui il nome del vendor su NVD differisce dall\'handle dell\'organizzazione GitHub (es. "nvidia" → "NVIDIA", "cisco" → "cisco-open-source")')
bullet('ricerca dentro l\'organizzazione del vendor con /search/repositories?q={product} in:name org:{vendor}')
p(
    'Una ricerca globale fuzzy come ulteriore fallback non viene utilizzata di '
    'proposito, perché tende a produrre falsi positivi (fork omonimi e progetti '
    'non correlati). Le coppie già risolte vengono memorizzate in una cache '
    'in-process: una stessa coppia che ricorre in molti CVE costa una sola '
    'chiamata API.'
)
p(
    'Il punteggio di selezione di un repository è il numero di CVE distinti che '
    'menzionano quella coppia (vendor, product) negli ultimi 30 giorni. L\'ipotesi '
    'di base: i prodotti pesantemente colpiti negli ultimi 30 giorni sono '
    'statisticamente i più probabili candidati a essere colpiti di nuovo nel '
    'prossimo periodo.'
)

# ----- Task 2 -----
h2('3.2 — Task "Hot": segnali di attività di sicurezza e silent patch')
p(
    'Questo task cerca repository che mostrano segnali di lavoro su problemi di '
    'sicurezza, anche quelli non dichiarati esplicitamente (silent patch). '
    'Procede in due fasi.'
)
p('Fase 1 — ricerca per parole chiave', bold=True)
p(
    'Per ciascuna delle parole chiave di sicurezza configurate '
    '(CVE, vulnerability, exploit, security, injection, XSS, CSRF, overflow, RCE, '
    'sanitize, auth bypass, credential, patch) interroghiamo l\'endpoint '
    '/search/commits di GitHub limitato agli ultimi 7 giorni, paginando fino a 3 '
    'pagine (300 commit) per parola chiave. I commit sono deduplicati per SHA, '
    'così che un singolo commit il cui messaggio combacia con più parole chiave '
    'venga contato una sola volta.'
)
p('Fase 2 — arricchimento con segnali di silent patch', bold=True)
p(
    'I repository candidati vengono arricchiti con due segnali aggiuntivi:'
)
bullet(
    'Numero di commit nell\'ultima settimana, ottenuto dall\'endpoint '
    '/repos/{owner}/{repo}/stats/commit_activity. Un picco improvviso di commit '
    'senza messaggi che parlano esplicitamente di sicurezza è un classico '
    'sintomo di silent patching.'
)
bullet(
    'Numero totale di download degli asset delle release, ottenuto sommando i '
    'campi download_count di /repos/{owner}/{repo}/releases. Questo è il proxy '
    'più affidabile disponibile su GitHub per stimare quanti utenti reali '
    'utilizzano il software (GitHub non espone un download count del codice '
    'sorgente). Più alto il download count, più ampia la superficie d\'impatto '
    'potenziale di una vulnerabilità.'
)
p('Lo score finale è una combinazione lineare additiva delle tre componenti:')
code('keyword_score   = #commit_unici + 2 × #parole_chiave_distinte\n'
     'commit_factor   = commits_ultima_settimana × 0,5\n'
     'download_factor = log10(download_totali + 1) × 3\n'
     '\n'
     'score = keyword_score + commit_factor + download_factor')
p(
    'I pesi (0,5 per i commit, 3 per il logaritmo dei download) sono stati scelti '
    'in modo che le tre componenti contribuiscano in modo paragonabile e sono '
    'modificabili in cima al file task_hot.py. La trasformazione logaritmica sui '
    'download evita che repository molto popolari (con download nell\'ordine dei '
    'milioni) saturino completamente lo score. Il numero di candidati arricchiti '
    'è limitato a MAX_REPOS_PER_TASK × 2 (parametro ENRICH_MULTIPLIER) per '
    'contenere il costo delle chiamate API.'
)

# ----- Task 3 -----
h2('3.3 — Task "Talkers": repository più rumorosi')
p(
    'Questo task identifica i repository che ricevono più "attenzione" su GitHub '
    'in questo momento, indipendentemente dal contenuto. L\'idea è che i repository '
    'più discussi, modificati e seguiti hanno una superficie d\'attacco più esposta '
    'e attirano più sguardi (compresi quelli degli attaccanti).'
)
p('Si interrogano due endpoint Search di GitHub limitati agli ultimi 7 giorni:')
bullet('/search/issues — issue create nel periodo')
bullet('/search/commits — commit autenticati nel periodo')
p(
    'I risultati vengono paginati fino a 1000 elementi per ciascun endpoint. Per '
    'ogni risultato si estrae il repository di appartenenza e si conta. Lo score '
    'è la somma pesata:'
)
code('score = W_ISSUES × #issue_recenti + W_COMMITS × #commit_recenti\n'
     '          (default 1.0)             (default 1.5)')
p(
    'Il maggior peso assegnato ai commit riflette il fatto che attività coordinata '
    'di sviluppo è un segnale più forte di esposizione rispetto al solo "rumore" '
    'degli utenti. Si selezionano i top 100 repository per score combinato.'
)

# =============================================================================
# 4. PERSISTENZA
# =============================================================================
h1('4. Persistenza dei dati')
p('Il database SQLite vulnRadar.db contiene tre tabelle principali:')

p('tabella tracked_repos', bold=True)
p(
    'Una riga per ogni selezione (repository, data, task). Include lo score '
    'numerico e una descrizione testuale del motivo della selezione (es. il numero '
    'di commit di sicurezza trovati). Il vincolo UNIQUE su '
    '(full_name, selected_date, task) garantisce che la stessa selezione non venga '
    'duplicata, ma il repository può apparire più volte nel tempo (in giorni '
    'diversi o selezionato da task diversi). Lo storico cresce ad ogni esecuzione.'
)

p('tabella cve_matches', bold=True)
p(
    'Registra ogni volta che un repository selezionato in passato compare in un '
    'nuovo CVE. Include il giorno di pubblicazione del CVE, il primo giorno in cui '
    'il repository era stato selezionato, e il numero di giorni intercorsi tra le '
    'due date (days_until_cve, sempre ≥ 0). Il vincolo UNIQUE su (repo, cve_id) '
    'garantisce che ogni match venga registrato una sola volta. Le righe non '
    'vengono mai cancellate né sovrascritte: il sistema accumula una storia '
    'permanente di "previsioni azzeccate".'
)

p('tabella last_check', bold=True)
p(
    'Tabella di servizio che memorizza l\'estremo superiore dell\'ultima '
    'finestra NVD scaricata con successo, in modo che ad ogni esecuzione il '
    'modulo di matching scarichi solo i CVE nuovi (approccio incrementale).'
)

# =============================================================================
# 5. CVE MATCHING
# =============================================================================
h1('5. Confronto con i CVE pubblicati successivamente')
p(
    'Dopo che i tre task hanno popolato tracked_repos con le selezioni del giorno, '
    'si esegue il modulo cve_matcher.py. Questo modulo:'
)
bullet('legge dalla tabella last_check il timestamp dell\'ultima esecuzione riuscita')
bullet('interroga il database NVD per tutti i CVE pubblicati da quel timestamp ad oggi tramite il client robusto')
bullet(
    'per ogni CVE estrae con un\'espressione regolare gli URL del tipo '
    'github.com/{owner}/{repo} dal campo references[]'
)
bullet(
    'incrocia ogni URL trovato con la lista completa di tutti i repository mai '
    'selezionati (storico di tracked_repos)'
)
bullet(
    'per ogni corrispondenza inserisce una riga in cve_matches con il numero di '
    'giorni di anticipo con cui il sistema aveva segnalato il repository'
)
bullet('aggiorna last_check con il timestamp coperto dall\'ultima finestra riuscita')

p(
    'Sono state implementate due garanzie metodologiche importanti per la '
    'validità scientifica dei risultati:'
)
p('Nessuna falsa previsione retroattiva', bold=True)
p(
    'I match vengono contati solo se il CVE è stato pubblicato lo stesso giorno '
    'o successivamente al giorno di prima selezione del repository. I CVE che '
    'precedono la selezione vengono scartati e contati separatamente nel log '
    '(skipped_pre_selection): non possono essere considerati previsioni.'
)
p('Nessuna perdita silenziosa di dati', bold=True)
p(
    'Il client NVD divide range temporali superiori a 119 giorni in finestre '
    'più piccole (NVD impone un limite massimo di 120 giorni per query) e '
    'restituisce esplicitamente l\'estremo superiore dell\'ultima finestra '
    'scaricata con successo. Il cursore last_check viene fatto avanzare solo '
    'fino a quel punto, mai a "now()". Se una finestra fallisce a metà, la '
    'coda mancante verrà ripresa automaticamente alla successiva esecuzione.'
)
p(
    'L\'estrattore di URL filtra anche i percorsi riservati GitHub '
    '(advisories/, orgs/, sponsors/, marketplace/, pulls/, issues/, …) che '
    'hanno la stessa forma di owner/repo ma non sono repository reali.'
)
p(
    'Al termine di ogni esecuzione il programma stampa la tabella delle '
    'corrispondenze più recenti (default 30), con il totale storico, in modo '
    'che siano sempre verificabili a colpo d\'occhio.'
)

# =============================================================================
# 6. PARAMETRI CONFIGURABILI
# =============================================================================
h1('6. Parametri configurabili')
p(
    'Tutti i parametri principali sono raggruppati in cima a config.py per essere '
    'modificati senza toccare la logica:'
)
code('MAX_REPOS_PER_TASK    = 100   # repo selezionati per task per esecuzione\n'
     'NVD_LOOKBACK_DAYS     = 30    # finestra temporale del task Official\n'
     'HOT_LOOKBACK_DAYS     = 7     # finestra temporale del task Hot\n'
     'TALKERS_LOOKBACK_DAYS = 7     # finestra temporale del task Talkers\n'
     'SECURITY_KEYWORDS     = [...]  # parole chiave usate dal task Hot')
p(
    'Pesi e parametri specifici dei singoli task:'
)
bullet('task_hot.py → W_COMMITS, W_DOWNLOADS, ENRICH_MULTIPLIER, SEARCH_PAGES_PER_KEYWORD')
bullet('task_talkers.py → W_ISSUES, W_COMMITS')

p('Credenziali (file .CVEfixes.ini)', bold=True)
p(
    'Il token GitHub è letto dallo stesso file .CVEfixes.ini usato dagli altri '
    'strumenti del progetto. È possibile aggiungere opzionalmente una API key '
    'NVD che velocizza le chiamate verso il database NVD di un fattore 10× '
    '(la pausa tra le pagine scende da 6 secondi a 0,6 secondi):'
)
code('[GitHub]\n'
     'token = ghp_xxxxxxxxxxxxxxxxxxxx\n'
     '\n'
     '[NVD]\n'
     'api_key = la-tua-chiave-nvd     # opzionale, richiedibile gratis')

# =============================================================================
# 7. SCENARIO D'USO E VALORE SCIENTIFICO
# =============================================================================
h1('7. Scenario d\'uso e valore scientifico')
p(
    'Il sistema è progettato per essere lanciato in modalità "daemon" e '
    'rimanere attivo perennemente. Al primo avvio esegue subito una '
    'pipeline completa, poi entra in un ciclo infinito che si risveglia '
    'ogni giorno all\'orario configurato in DAILY_RUN_HOUR_UTC (default '
    'le 6:00 UTC) e ripete il lavoro. Eventuali errori in una singola '
    'esecuzione vengono catturati e registrati nei log, senza '
    'interrompere il ciclo: una transitoria indisponibilità delle API '
    'di GitHub o NVD non blocca le esecuzioni successive. Per terminare '
    'il programma è sufficiente un Ctrl+C, che viene gestito in modo '
    'pulito sia durante l\'esecuzione che durante l\'attesa.'
)
p(
    'In alternativa è possibile impostare DAILY_RUN_HOUR_UTC = None per '
    'eseguire una singola pipeline e terminare, lasciando la '
    'schedulazione a strumenti esterni come cron o systemd.'
)
p(
    'Ogni esecuzione richiede pochi minuti grazie all\'esecuzione '
    'parallela dei task e all\'uso opzionale della API key NVD. Nel '
    'tempo, il database accumula naturalmente uno storico che permette '
    'analisi a posteriori:'
)
bullet(
    'qual è il "tasso di hit" di ciascuno dei tre criteri? (numero di '
    'repository selezionati che effettivamente sono finiti in un CVE, '
    'diviso il totale selezionato)'
)
bullet(
    'qual è il tempo medio di anticipo (days_until_cve) — ovvero quanti giorni '
    'prima della pubblicazione ufficiale del CVE il sistema aveva segnalato il '
    'repository?'
)
bullet(
    'esiste un task strutturalmente migliore degli altri? esistono combinazioni '
    'di segnali (es. repository selezionato da due task contemporaneamente) che '
    'predicono meglio?'
)
p(
    'Queste sono le domande di ricerca che il database risponderà man mano che '
    'l\'osservazione prosegue. Le garanzie metodologiche introdotte (nessuna '
    'falsa previsione retroattiva, nessuna perdita silenziosa di dati) '
    'garantiscono che i tassi di hit calcolati siano statisticamente onesti.'
)

# =============================================================================
# 8. BIAS NOTI E LIMITAZIONI METODOLOGICHE
# =============================================================================
h1('8. Bias noti e limitazioni metodologiche')
p(
    'Per onestà scientifica vengono qui dichiarati i principali bias che '
    'possono influenzare i risultati prodotti dal sistema, in modo che il '
    'lettore possa interpretare correttamente le statistiche raccolte.'
)

p('Bias di campionamento del Task 3 (Talkers)', bold=True)
p(
    'L\'API di ricerca di GitHub (/search/issues e /search/commits) restituisce '
    'al massimo 1000 risultati per query, indipendentemente dal numero totale '
    'di elementi che soddisfano i criteri. Su finestre temporali ampie come 7 '
    'giorni e con criteri molto generici come "tutti i commit globali", su '
    'GitHub vengono creati centinaia di commit e issue ogni minuto: 1000 '
    'risultati ordinati per data decrescente coprono in pratica solo le ultime '
    'ore (a volte i soli ultimi minuti) del periodo richiesto.'
)
p(
    'La conseguenza è che il Task 3 non identifica veramente "i repository più '
    'attivi degli ultimi 7 giorni" ma piuttosto "i repository più attivi nel '
    'momento esatto in cui il programma viene eseguito". I repository che '
    'fanno burst di attività distribuiti nell\'arco della settimana ma più '
    'silenziosi nelle ore precedenti l\'esecuzione vengono sottostimati.'
)
p(
    'Una mitigazione possibile sarebbe paginare per fasce temporali più '
    'piccole (ad esempio scaricare separatamente ciascun giorno della '
    'finestra di 7 giorni) e poi unire i risultati. Questo aumenterebbe '
    'significativamente il costo in chiamate API ma rimuoverebbe quasi '
    'completamente il bias. La modifica non è stata fatta in questa versione '
    'per mantenere il sistema veloce e prevedibile sotto i limiti di rate '
    'imposti da GitHub; è documentata come miglioria futura. Quando si '
    'analizzano i risultati del Task 3, è importante ricordare che '
    'l\'orario di esecuzione del cron job condiziona l\'output: due '
    'esecuzioni a 12 ore di distanza producono insiemi di "talkers" '
    'parzialmente diversi.'
)

p('Bias di copertura del Task 1 (Official)', bold=True)
p(
    'Il task assume che i prodotti più colpiti negli ultimi 30 giorni siano '
    'rappresentativi anche del prossimo periodo. Questa ipotesi vale sui '
    'tempi medi ma può perdere prodotti completamente nuovi che vengono '
    'colpiti per la prima volta. Inoltre, NVD pubblica i CVE con un certo '
    'ritardo rispetto alla scoperta delle vulnerabilità: i 30 giorni '
    'osservati non corrispondono ai 30 giorni di reali ricerche di '
    'sicurezza in corso.'
)

p('Bias di lessico del Task 2 (Hot)', bold=True)
p(
    'La ricerca per parole chiave premia i progetti la cui cultura di '
    'commit è esplicita su questioni di sicurezza (ad esempio progetti '
    'che includono "fix CVE-XXXX" nel messaggio). I progetti con '
    'convenzioni di commit minimaliste o con lingua diversa dall\'inglese '
    'sono sistematicamente sottorappresentati nel Task 2, anche quando '
    'fanno fix di sicurezza importanti. I segnali di silent patch '
    '(commit_factor e download_factor) compensano parzialmente questo '
    'bias ma non lo eliminano.'
)

p('Sintesi', bold=True)
p(
    'Le statistiche di hit-rate calcolate confrontando le selezioni con i '
    'CVE pubblicati successivamente sono interpretabili in senso relativo '
    '(quale criterio fra i tre ha la migliore precisione predittiva, '
    'come evolvono i tassi nel tempo) più che in senso assoluto. I bias '
    'qui dichiarati non invalidano le previsioni positive (un CVE che il '
    'sistema ha previsto è stato veramente previsto), ma indicano che '
    'esistono CVE che il sistema non avrebbe potuto prevedere per ragioni '
    'strutturali e che quindi non vanno conteggiati come "errori del '
    'modello" nelle analisi di precisione.'
)

# =============================================================================
# SAVE
# =============================================================================
output_path = '/home/medo/laMiaTesi/vulnRadar/vulnRadar_spiegazione.docx'
doc.save(output_path)
print(f'Documento creato: {output_path}')
